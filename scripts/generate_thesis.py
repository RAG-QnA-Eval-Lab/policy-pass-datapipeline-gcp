#!/usr/bin/env python3
"""졸업논문 .docx 생성 스크립트.

Usage:
    python scripts/generate_thesis.py

Output:
    docs/졸업논문_김대현_2026_완성.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

TEMPLATE = Path("docs/졸업논문_김대현_2026.docx")
OUTPUT = Path("docs/졸업논문_김대현_2026_완성.docx")

S_TITLE = "MDPI_1.2_title"
S_AUTHOR = "MDPI_1.3_authornames"
S_AFFIL = "MDPI_1.6_affiliation"
S_ABSTRACT = "MDPI_1.7_abstract"
S_KEYWORDS = "MDPI_1.8_keywords"
S_LINE = "MDPI_1.9_line"
S_H1 = "MDPI_2.1_heading1"
S_H2 = "MDPI_2.2_heading2"
S_H3 = "MDPI_2.3_heading3"
S_TEXT = "MDPI_3.1_text"
S_ITEM = "MDPI_3.7_itemize"
S_BULLET = "MDPI_3.8_bullet"
S_TABLE_CAP = "MDPI_4.1_table_caption"
S_TABLE_FOOT = "MDPI_4.3_table_footer"
S_FIG_CAP = "MDPI_5.1_figure_caption"
S_FIG = "MDPI_5.2_figure"
S_BACK = "MDPI_6.2_BackMatter"
S_REF = "MDPI_7.1_References"


class ThesisWriter:
    def __init__(self):
        self.doc = Document(str(TEMPLATE))
        self._clear_body()

    # ── infrastructure ─────────────────────────────────
    def _clear_body(self):
        body = self.doc.element.body
        final_sect = None
        for child in reversed(list(body)):
            if child.tag.endswith("}sectPr"):
                final_sect = child
                break
        for child in list(body):
            body.remove(child)
        if final_sect is not None:
            body.append(final_sect)

    def p(self, text: str, style: str = S_TEXT):
        para = self.doc.add_paragraph(text)
        try:
            para.style = self.doc.styles[style]
        except KeyError:
            pass
        return para

    def h1(self, text: str):
        return self.p(text, S_H1)

    def h2(self, text: str):
        return self.p(text, S_H2)

    def h3(self, text: str):
        return self.p(text, S_H3)

    def bullet(self, text: str):
        return self.p(text, S_BULLET)

    def item(self, text: str):
        return self.p(text, S_ITEM)

    def fig_placeholder(self, caption: str):
        self.p("[그림 삽입 위치]", S_FIG)
        self.p(caption, S_FIG_CAP)

    def table(self, headers: list[str], rows: list[list[str]], caption: str | None = None, footer: str | None = None):
        if caption:
            self.p(caption, S_TABLE_CAP)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                tbl.rows[r_idx + 1].cells[c_idx].text = str(val)
        if footer:
            self.p(footer, S_TABLE_FOOT)
        return tbl

    def equation(self, formula: str, number: str):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.columns[0].width = Cm(12)
        tbl.columns[1].width = Cm(2)
        tbl.rows[0].cells[0].text = formula
        tbl.rows[0].cells[1].text = number
        tbl.rows[0].cells[0].paragraphs[0].alignment = 1  # center
        tbl.rows[0].cells[1].paragraphs[0].alignment = 2  # right

    def page_break(self):
        para = self.doc.add_paragraph()
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._element.append(br)

    def section_break(self):
        self.doc.add_section(WD_ORIENT.PORTRAIT)

    # ── cover page ─────────────────────────────────────
    def write_cover(self):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.text = (
            "본 논문은 2026학년 2월 졸업을 위해 제출된 "
            "한국외국어대학교 컴퓨터공학부 졸업논문이다. 2026.02.\n"
            "지도교수: 장익범\n"
            "서명: ______________________\n"
        )
        self.p("", S_LINE)
        self.section_break()

    # ── header ─────────────────────────────────────────
    def write_header(self):
        self.p(
            "RAG 응답 모니터링을 위한 비용 효율적 LLM-as-a-Judge "
            "파이프라인 구축 및 성능 비교 분석",
            S_TITLE,
        )
        self.p(
            "Cost-Efficient LLM-as-a-Judge Pipeline for RAG Response "
            "Monitoring: Design, Implementation, and Comparative Analysis",
            S_TITLE,
        )
        self.p("김대현 (Daehyun Kim) 1", S_AUTHOR)
        self.p(
            "1 한국외국어대학교, 컴퓨터공학부; daehyun@hufs.ac.kr",
            S_AFFIL,
        )

    # ── abstract ───────────────────────────────────────
    def write_abstract(self):
        self.p("", S_LINE)
        self.p(
            "한글 요약: "
            "본 연구는 대한민국 청년·학생 대상 정부 정책 정보의 접근성 개선을 목적으로, "
            "하이브리드 RAG(Retrieval-Augmented Generation) 기반 질의응답 시스템과 "
            "비용 효율적 3단계 응답 품질 평가 파이프라인을 설계·구현하였다. "
            "데이터 수집 단계에서는 공공데이터포털 및 온통청년 API를 통해 2,235건의 "
            "청년 정책 문서를 수집하고, 한국어 문장 경계 기반 청킹(512토큰, 50토큰 오버랩)을 "
            "적용하였다. 검색 단계에서는 FAISS 기반 Dense Retrieval과 BM25 Sparse Retrieval을 "
            "Reciprocal Rank Fusion(k=60)으로 결합한 후 Cross-Encoder 리랭킹을 수행하는 "
            "4단계 하이브리드 파이프라인을 구축하였다. 생성 단계에서는 LiteLLM을 통해 GPT-4o, "
            "GPT-4o-mini, Claude Sonnet, Gemini Flash, Llama3 다섯 모델을 통합하고, "
            "RAG 적용 유무에 따른 응답 품질 변화를 비교하였다. 평가 단계에서는 RAGAS v0.4 "
            "정량 지표, G-Eval 방식의 LLM Judge, DeepEval 환각 탐지를 결합한 3단계 "
            "파이프라인을 구현하였으며, Position Bias 완화를 위해 컨텍스트 순서를 두 차례 "
            "달리한 평균 점수를 사용하였다. 실험 결과, 하이브리드 리랭킹 전략이 BM25 단독 "
            "검색 대비 Context Recall 기준 16.6% 높은 0.855를 달성하였으며, GPT-4o-mini를 "
            "Judge 모델로 활용한 경우 Gemini 3.1 Pro 대비 90.3% 수준의 평가 일관성을 "
            "입력 토큰 기준 약 13배, 출력 토큰 기준 약 20배 낮은 비용으로 달성함을 확인하였다.",
            S_ABSTRACT,
        )
        self.p(
            "핵심어: 검색 증강 생성, LLM-as-a-Judge, 하이브리드 검색, 자동 평가, 청년 정책",
            S_ABSTRACT,
        )
        self.p(
            "영문 요약: "
            "This study designs and implements a hybrid Retrieval-Augmented Generation (RAG) "
            "question-answering system for Korean government youth policy information, "
            "accompanied by a cost-efficient three-stage response quality evaluation pipeline. "
            "We collected 2,235 youth policy documents via the Public Data Portal and Youth "
            "Policy APIs, and applied Korean sentence-boundary-aware chunking (512 tokens, "
            "50-token overlap). For retrieval, we constructed a four-strategy hybrid pipeline "
            "combining FAISS-based dense retrieval with BM25 sparse retrieval through "
            "Reciprocal Rank Fusion (k=60) and Cross-Encoder reranking. Five large language "
            "models — GPT-4o, GPT-4o-mini, Claude Sonnet, Gemini Flash, and Llama3 — were "
            "unified under LiteLLM and compared with and without RAG augmentation. Evaluation "
            "employed a three-stage pipeline: (1) RAGAS v0.4 reference-based metrics "
            "(Faithfulness, Answer Relevancy, Context Precision, Context Recall); (2) "
            "G-Eval-style LLM-as-a-Judge scoring three dimensions (citation accuracy, "
            "completeness, readability) with position bias mitigation via two-pass context "
            "shuffling; and (3) DeepEval hallucination detection. Experiments on 100 curated "
            "QA pairs show that the hybrid reranking strategy achieves 16.6% higher Context "
            "Recall than BM25-only retrieval, and that GPT-4o-mini as the judge model attains "
            "90.3% class agreement with Gemini 3.1 Pro at approximately 13x lower input token cost, "
            "demonstrating the viability of cost-efficient LLM-as-a-Judge monitoring for "
            "production RAG systems.",
            S_ABSTRACT,
        )
        self.p(
            "Keywords: Retrieval-Augmented Generation, LLM-as-a-Judge, Hybrid Retrieval, "
            "Automated Evaluation, Youth Policy",
            S_KEYWORDS,
        )
        self.p("", S_LINE)

    # ── 1. 서론 ────────────────────────────────────────
    def write_section1(self):
        self.h1("1. 서론")

        self.h2("1.1. 연구 배경 및 동기")
        self.p(
            "대한민국 정부는 매년 수백 건의 청년·학생 대상 지원 정책을 신설하거나 변경한다. "
            "고용노동부, 중소벤처기업부, 국토교통부 등 다수의 부처가 취업, 창업, 주거, 교육, "
            "복지, 금융 등 다양한 분야에서 정책을 시행하고 있으나, 이러한 정보는 각 부처 웹사이트에 "
            "분산되어 있어 청년 수혜자가 자신에게 적합한 정책을 탐색하는 데 상당한 시간과 노력이 "
            "소요된다. 온통청년(www.youthcenter.go.kr)과 같은 통합 포털이 존재하나, 키워드 기반 "
            "검색만 지원하며 자연어 질의에 대한 맥락적 응답 기능은 제공하지 않는다."
        )
        self.p(
            "최근 대형 언어 모델(LLM, Large Language Model)의 발전으로 자연어 질의응답 시스템의 "
            "구축 비용이 크게 낮아졌다. 그러나 LLM은 학습 데이터에 포함되지 않은 최신 정보에 대해 "
            "사실과 다른 내용을 생성하는 환각(Hallucination) 문제를 내재하고 있다 [1]. "
            "특히 정부 정책 정보처럼 수시로 갱신되는 도메인에서는 이러한 문제가 더욱 심각하다. "
            "Lewis 등 [1]이 제안한 RAG(Retrieval-Augmented Generation)는 외부 문서를 검색하여 "
            "LLM의 생성에 근거를 제공함으로써 이 문제를 완화하지만, RAG 시스템 자체의 응답 품질을 "
            "체계적으로 평가하고 모니터링하는 과제가 남아 있다."
        )
        self.p(
            "RAG 시스템의 응답 품질 평가에는 크게 두 가지 접근법이 존재한다. 첫째, RAGAS [2]와 "
            "같은 참조 기반 자동 평가 프레임워크는 Faithfulness, Context Precision 등의 정량 "
            "지표를 제공하지만, 답변의 가독성이나 인용 정확도와 같은 정성적 품질은 포착하지 못한다. "
            "둘째, LLM-as-a-Judge [3] 패러다임은 LLM 자체를 평가자로 활용하여 정성적 차원을 "
            "측정할 수 있으나, 고성능 Judge 모델(GPT-4o, Gemini Pro 등)의 평가 비용이 생성 비용을 "
            "초과할 수 있다는 실무적 한계가 있다. 이러한 비용 문제는 특히 프로덕션 환경에서 "
            "지속적 모니터링을 수행해야 하는 경우 치명적이다."
        )

        self.h2("1.2. 연구 목적 및 범위")
        self.p("본 연구의 목적은 다음과 같다.")
        self.item(
            "1. 한국어 청년 정책 도메인에 특화된 하이브리드 RAG 질의응답 시스템을 구축하고, "
            "4가지 검색 전략(Vector, BM25, Hybrid RRF, Hybrid Rerank) 간 성능을 체계적으로 비교한다."
        )
        self.item(
            "2. RAGAS v0.4 정량 평가, G-Eval 방식 LLM-as-a-Judge 정성 평가, DeepEval 환각 "
            "탐지를 결합한 3단계 자동 평가 파이프라인을 설계하고, 각 단계의 상보성을 실증한다."
        )
        self.item(
            "3. 파인튜닝 없이 경량 상용 모델(GPT-4o-mini)을 Judge로 활용하는 방안의 비용 "
            "효율성을 고비용 모델(Gemini 3.1 Pro) 대비 검증한다."
        )
        self.item(
            "4. RAG 컨텍스트 순서에 의한 Position Bias를 완화하기 위한 2회 평균 기법의 "
            "효과를 실증한다."
        )
        self.p(
            "연구 범위는 한국어 텍스트 기반 정부 정책 문서에 한정하며, 100쌍의 QA 데이터셋을 "
            "사용하여 5개 LLM(GPT-4o, GPT-4o-mini, Claude Sonnet, Gemini Flash, Llama3)의 "
            "응답 품질을 비교한다. 시스템은 Google Cloud Platform(GCP) Cloud Run에 배포하여 "
            "서버리스 환경에서의 운영 가능성을 검증한다."
        )

        self.h2("1.3. 논문 구성")
        self.p(
            "본 논문의 나머지 부분은 다음과 같이 구성된다. 2장에서는 RAG 시스템, LLM 평가 "
            "방법론, 비용 효율적 평가 연구에 대한 관련 연구를 고찰한다. 3장에서는 전체 시스템 "
            "아키텍처와 각 파이프라인의 설계 및 구현을 상세히 기술한다. 4장에서는 실험 설정과 "
            "결과를 분석하고, 5장에서 결론 및 향후 연구 방향을 제시한다."
        )

    # ── 2. 관련 연구 ───────────────────────────────────
    def write_section2(self):
        self.h1("2. 관련 연구")

        self.h2("2.1. RAG 기반 시스템")

        self.h3("2.1.1. Hybrid 검색과 리랭킹")
        self.p(
            "Lewis 등 [1]이 제안한 RAG는 외부 지식 베이스에서 관련 문서를 검색하여 LLM의 "
            "생성에 근거를 제공하는 기법이다. 초기 RAG 시스템은 단일 Dense Retriever에 "
            "의존하였으나, 최근에는 Dense Retrieval과 Sparse Retrieval을 결합한 Hybrid 검색이 "
            "성능 향상에 효과적임이 보고되고 있다."
        )
        self.p(
            "Cormack 등 [8]은 여러 검색 시스템의 결과를 통합하는 Reciprocal Rank "
            "Fusion(RRF) 기법을 제안하였다. RRF는 각 검색기의 순위를 역수 변환하여 "
            "합산함으로써, 단일 검색기 대비 안정적인 성능 향상을 달성한다. 또한 Nogueira와 "
            "Cho [9]는 BERT 기반 Cross-Encoder를 통한 passage reranking이 검색 정밀도를 "
            "크게 향상시킴을 보였다."
        )
        self.p(
            "그러나 기존 연구 대부분은 영어 도메인에 집중되어 있으며, 한국어 정책 문서와 같은 "
            "특수 도메인에서의 Hybrid 검색 전략 간 체계적 비교는 부족한 실정이다. 본 연구는 "
            "이러한 공백을 메우기 위해 4가지 검색 전략을 한국어 정책 도메인에서 실증적으로 비교한다."
        )

        self.h3("2.1.2. 청크 전략과 한국어 처리")
        self.p(
            "RAG 시스템의 성능은 문서 청킹(Chunking) 전략에 크게 영향을 받는다. 고정 크기 "
            "청킹은 구현이 단순하나 의미 단위를 훼손할 수 있으며, 문장 경계 기반 청킹은 의미 "
            "보존에 유리하지만 언어별 문장 분리기가 필요하다. 한국어의 경우 영어와 달리 문장 "
            "경계 인식이 복잡하여, 전용 한국어 문장 분리기(kss 등)의 활용이 필수적이다. "
            "본 연구에서는 kss를 이용한 문장 경계 인식 후 tiktoken cl100k_base 토크나이저로 "
            "512토큰 단위 청킹을 수행하며, 50토큰 오버랩을 적용하여 문맥 손실을 방지한다."
        )

        self.h2("2.2. LLM 평가 방법론")

        self.h3("2.2.1. 참조 기반 자동 평가: RAGAS")
        self.p(
            "Es 등 [2]은 RAG 시스템 전용 평가 프레임워크인 RAGAS를 제안하였다. RAGAS는 "
            "Faithfulness(답변이 컨텍스트에 근거하는 정도), Answer Relevancy(답변이 질문에 "
            "부합하는 정도), Context Precision(검색된 컨텍스트의 정밀도), Context Recall"
            "(검색된 컨텍스트의 재현율) 4가지 지표를 LLM 기반으로 자동 산출한다."
        )
        self.p(
            "본 연구에서는 RAGAS v0.4 API를 사용한다. v0.3과 v0.4는 API 인터페이스가 "
            "상이하여, v0.3의 evaluate() 일괄 평가 대신 v0.4의 SingleTurnSample 기반 "
            "metric.single_turn_ascore() 개별 평가를 사용한다. 이 버전 차이는 재현성에 "
            "중요하므로 명시한다."
        )

        self.h3("2.2.2. LLM-as-a-Judge 패러다임")
        self.p(
            "Zheng 등 [3]은 강력한 LLM을 평가자(Judge)로 활용하는 LLM-as-a-Judge 패러다임을 "
            "제안하였다. MT-Bench를 통한 실험에서 GPT-4의 평가가 인간 전문가의 평가와 80% 이상의 "
            "일치율을 보이며, LLM이 효과적인 자동 평가자로 기능할 수 있음을 실증하였다."
        )
        self.p(
            "Liu 등 [4]은 G-Eval 프레임워크에서 GPT-4를 활용한 NLG 평가가 기존 자동 평가 "
            "지표 대비 인간 판단과의 상관관계가 높음을 보였다. G-Eval은 평가 기준과 채점 근거를 "
            "체계적으로 구성한 프롬프트를 통해 세분화된 정성 평가를 수행한다."
        )
        self.p(
            "본 연구는 G-Eval 방식을 채택하여 인용 정확도(Citation Accuracy), 완결성"
            "(Completeness), 가독성(Readability) 3가지 차원에서 1~5점 정수 척도로 평가한다. "
            "이는 RAGAS가 측정하지 못하는 정성적 품질 차원을 보완하기 위함이다."
        )

        self.h3("2.2.3. Position Bias 문제와 완화 기법")
        self.p(
            "Wang 등 [5]은 LLM-as-a-Judge에서 평가 대상의 제시 순서가 점수에 영향을 미치는 "
            "Position Bias를 체계적으로 분석하였다. 이들의 연구는 주로 pairwise 비교(두 답변 "
            "A, B의 순서 교환)에서의 bias를 다루었다."
        )
        self.p(
            "본 연구에서는 pairwise가 아닌 pointwise 평가 환경에서, 검색된 컨텍스트 문서의 "
            "제시 순서가 Judge 점수에 미치는 영향을 분석한다. 이를 완화하기 위해 원본 순서 1회와 "
            "셔플 순서 1회, 총 2회 평가의 평균 점수를 사용하는 기법을 적용한다. 이는 Wang 등의 "
            "답변 순서 bias와는 다른 종류의 position bias로, RAG 특화 평가에서의 새로운 "
            "관점을 제시한다."
        )

        self.h2("2.3. 비용 효율적 LLM 평가 연구")
        self.p(
            "Kim 등 [6]은 오픈소스 LLM을 평가 목적으로 파인튜닝한 Prometheus를 제안하여, "
            "상용 모델에 의존하지 않는 평가 방안을 탐구하였다. Saad-Falcon 등 [7]은 소량의 "
            "라벨 데이터로 분류기를 학습시켜 RAG 시스템을 평가하는 ARES 프레임워크를 제안하였다."
        )
        self.p(
            "그러나 이들 접근법은 파인튜닝이나 학습 데이터 준비에 추가 비용과 전문 지식이 "
            "필요하다. 본 연구는 별도의 파인튜닝 없이 기존 상용 경량 모델(GPT-4o-mini)을 "
            "Judge로 직접 활용하여, 고비용 모델 대비 동등한 평가 품질을 달성할 수 있는지를 "
            "실증한다. 이는 파인튜닝 인프라를 갖추지 못한 실무 환경에서의 실용적 기여를 "
            "목표로 한다."
        )

    # ── 3. 시스템 설계 및 구현 ──────────────────────────
    def write_section3(self):
        self.h1("3. 시스템 설계 및 구현")

        # 3.1
        self.h2("3.1. 전체 시스템 아키텍처")
        self.p(
            "본 시스템은 데이터 수집, 인덱싱, 검색, 생성, 평가의 5단계 파이프라인으로 "
            "구성된다(그림 1). 전체 인프라는 Google Cloud Platform(GCP) 위에 구축되며, "
            "Cloud Storage(GCS)를 데이터 저장소로, Compute Engine VM을 MongoDB 메타데이터 "
            "서버와 Airflow 스케줄러로, Cloud Run을 서빙 환경으로 활용한다."
        )
        self.p(
            "FastAPI 기반 백엔드는 6개 REST API 엔드포인트(Health, Search, Generate, "
            "Policies, Models, Evaluate)를 제공하며, Cloud Run의 scale-to-zero 특성을 "
            "활용하여 비용을 최소화한다. 프론트엔드는 Streamlit으로 구현하여 별도의 Cloud Run "
            "인스턴스에 배포한다."
        )
        self.fig_placeholder("그림 1. 전체 시스템 아키텍처.")

        # 3.2
        self.h2("3.2. 데이터 수집 및 전처리")

        self.h3("3.2.1. 공공 정책 데이터 수집 파이프라인")
        self.p(
            "본 시스템은 공공데이터포털(data.go.kr) API와 온통청년(youthcenter.go.kr) API를 "
            "통해 청년 정책 데이터를 수집한다(표 1). 각 데이터 소스별로 전용 수집기(Collector)를 "
            "구현하여, 수집된 원본 데이터를 Policy frozen dataclass로 정규화한다. 정규화된 데이터는 "
            "GCS에 원본으로 저장되고, MongoDB에 메타데이터가 기록된다."
        )
        self.table(
            ["소스", "수집 방식", "수집 건수", "주요 카테고리"],
            [
                ["공공데이터포털", "REST API", "~1,500", "취업, 창업, 주거, 금융"],
                ["온통청년", "REST API", "~735", "교육, 복지, 주거"],
                ["합계", "-", "2,235", "6개 카테고리"],
            ],
            caption="표 1. 데이터 소스 요약.",
        )
        self.p(
            "Airflow DAG을 통해 수집과 인덱싱을 매일 02:00(KST)에 자동 실행하며, "
            "robots.txt를 준수하고 요청 간격 2~3초를 유지하여 서버 부하를 방지한다. "
            "총 2,235건의 청년 정책 문서를 수집하였으며, 취업, 주거, 교육, 복지 등의 "
            "카테고리로 분류된다."
        )

        self.h3("3.2.2. 한국어 청킹 전략")
        self.p(
            "수집된 문서는 한국어 문장 경계 기반 청킹을 거쳐 검색 가능한 단위로 분할된다. "
            "kss(Korean Sentence Splitter) 라이브러리로 문장을 분리한 후, tiktoken "
            "cl100k_base 토크나이저를 사용하여 512토큰 단위로 청크를 구성한다. 인접 청크 간 "
            "50토큰의 오버랩을 적용하여 문맥 단절을 방지한다. kss가 설치되지 않은 환경에서는 "
            "mecab, 구두점 기반, regex 순으로 폴백한다."
        )
        self.p(
            "이 방식은 고정 크기 청킹 대비 의미 단위를 보존하는 장점이 있다. 특히 정부 정책 "
            "문서는 조항별 구조가 명확하여 문장 경계 기반 분할이 효과적이다."
        )

        # 3.3
        self.h2("3.3. 하이브리드 검색 파이프라인")
        self.p(
            "본 시스템은 4가지 검색 전략을 SearchStrategy enum으로 관리한다(그림 2). "
            "사용자는 API 호출 시 전략을 선택할 수 있으며, 기본값은 hybrid_rerank이다."
        )
        self.fig_placeholder("그림 2. 하이브리드 검색 파이프라인 흐름도.")

        self.h3("3.3.1. Dense Retrieval (FAISS)")
        self.p(
            "FAISS(Facebook AI Similarity Search)를 이용한 Dense Retrieval을 기본 검색기로 "
            "사용한다. LiteLLM embedding API를 통해 text-embedding-3-small 모델로 문서 "
            "임베딩을 생성하고, FAISS IndexFlatIP 인덱스에 저장한다. FAISS를 선택한 이유는 "
            "Cloud Run scale-to-zero 환경에서 인메모리 인덱스가 적합하기 때문이다. ChromaDB와 "
            "같은 클라이언트-서버 구조는 별도 프로세스가 필요하여 서버리스 환경에 부적합하다."
        )

        self.h3("3.3.2. BM25 Sparse Retrieval")
        self.p(
            "BM25 알고리즘을 이용한 키워드 기반 검색을 보조 검색기로 구현한다. 한국어 "
            "토크나이징은 공백 기반 분리를 사용하며, 이는 형태소 분석기 의존성을 줄이면서도 "
            "정책 문서의 고유명사와 전문 용어 매칭에 효과적이다."
        )

        self.h3("3.3.3. Reciprocal Rank Fusion")
        self.p(
            "Dense Retrieval과 BM25의 결과를 Reciprocal Rank Fusion(RRF) [8]으로 통합한다. "
            "RRF 점수는 수식 (1)과 같이 계산된다."
        )
        self.equation("RRF(d) = Σ_{r∈R}  1 / (k + rank_r(d))", "(1)")
        self.p(
            "여기서 k는 상수(본 연구에서 k=60), rank_r(d)는 검색기 r에서 문서 d의 순위, "
            "R은 검색기 집합이다. k=60은 Cormack 등 [8]의 권장값으로, 상위 순위에 과도한 "
            "가중치가 집중되는 것을 방지한다."
        )

        self.h3("3.3.4. Cross-Encoder 리랭킹")
        self.p(
            "RRF 통합 결과에 대해 Cross-Encoder 모델(ms-marco-MiniLM-L-6-v2)로 리랭킹을 "
            "수행한다. Cross-Encoder는 질의-문서 쌍을 동시에 입력받아 관련성 점수를 산출하므로, "
            "Bi-Encoder 기반 Dense Retrieval보다 정밀한 관련성 판단이 가능하다 [9]. "
            "ms-marco-MiniLM-L-6-v2를 선택한 이유는 Cloud Run 2GiB 메모리 제약 하에서 "
            "로드 가능한 경량 모델이기 때문이다."
        )

        # 3.4
        self.h2("3.4. 멀티 LLM 생성 파이프라인")
        self.p(
            "LiteLLM을 통해 5개 LLM을 통합한다(표 2). LiteLLM은 단일 API 인터페이스로 "
            "OpenAI, Vertex AI(Google), HuggingFace 등 다양한 프로바이더의 모델을 호출할 수 "
            "있어, 멀티 모델 비교 실험에 적합하다."
        )
        self.table(
            ["모델", "LiteLLM ID", "프로바이더"],
            [
                ["GPT-4o", "openai/gpt-4o", "OpenAI"],
                ["GPT-4o-mini", "openai/gpt-4o-mini", "OpenAI"],
                ["Claude Sonnet", "vertex_ai/claude-sonnet-4-5", "Vertex AI"],
                ["Gemini Flash", "vertex_ai/gemini-2.5-flash", "Vertex AI"],
                ["Llama3 (70B)", "huggingface/meta-llama/Llama-3.3-70B-Instruct", "HuggingFace"],
            ],
            caption="표 2. 사용 모델 요약.",
        )
        self.p(
            "모델 키와 LiteLLM ID 간의 매핑은 설정 파일에서 관리하며, Vertex AI 모델의 경우 "
            "리전별 가용성이 다르므로 모델별 리전 오버라이드를 적용한다. 생성 시에는 검색된 "
            "top-k 문서를 프롬프트에 포함하는 RAG 모드와, 컨텍스트 없이 질문만 전달하는 "
            "No-RAG 모드를 지원한다."
        )

        # 3.5
        self.h2("3.5. 3단계 평가 파이프라인")
        self.p(
            "본 연구의 핵심 기여인 3단계 평가 파이프라인은 각 단계가 서로 다른 RAG 실패 모드를 "
            "탐지하도록 설계되었다(그림 3)."
        )
        self.fig_placeholder("그림 3. 3단계 평가 파이프라인 흐름도.")

        self.h3("3.5.1. Stage 1: RAGAS 정량 평가")
        self.p(
            "RAGAS v0.4의 SingleTurnSample 기반으로 4개 지표를 산출한다. Faithfulness은 "
            "답변의 각 주장이 컨텍스트에 근거하는 비율을, Answer Relevancy는 답변이 질문 의도에 "
            "부합하는 정도를 측정한다. Context Precision과 Context Recall은 검색 단계의 품질을 "
            "평가한다. 이 단계는 주로 근거 부족(unsupported claims) 실패 모드를 탐지한다."
        )

        self.h3("3.5.2. Stage 2: LLM-as-a-Judge 정성 평가")
        self.p(
            "G-Eval [4] 방식으로 LLM을 Judge로 활용하여 3가지 차원을 1~5점 정수 척도로 "
            "평가한다."
        )
        self.bullet("인용 정확도(Citation Accuracy): 답변이 검색된 컨텍스트를 정확히 인용하는가;")
        self.bullet("완결성(Completeness): 질문에 대한 답변이 충분히 포괄적인가;")
        self.bullet("가독성(Readability): 답변의 구조와 표현이 이해하기 쉬운가.")
        self.p(
            "Position Bias 완화를 위해 원본 컨텍스트 순서로 1회, 셔플된 순서로 1회, "
            "총 2회 평가하여 평균 점수를 사용한다. 이 단계는 부정확한 인용 및 왜곡 실패 "
            "모드를 탐지한다."
        )

        self.h3("3.5.3. Stage 3: DeepEval 안전성 평가")
        self.p(
            "DeepEval 프레임워크의 HallucinationMetric을 사용하여 답변이 컨텍스트와 명시적으로 "
            "모순되는지를 탐지한다. 0.0(환각 없음)~1.0(완전한 환각) 범위의 점수를 산출한다. "
            "이 단계는 RAGAS Faithfulness와 유사한 개념이나, 명시적 모순(contradiction) 탐지에 "
            "특화되어 있어 상보적 역할을 수행한다."
        )

        # 3.6
        self.h2("3.6. GCP 배포 및 운영 인프라")
        self.table(
            ["구성 요소", "GCP 서비스", "사양", "용도"],
            [
                ["BE 서버", "Cloud Run", "2GiB, max 1", "FastAPI + FAISS 인메모리 검색"],
                ["FE 서버", "Cloud Run", "512MiB, max 1", "Streamlit UI"],
                ["데이터 저장소", "Cloud Storage", "-", "원본 문서 + FAISS 인덱스"],
                ["메타데이터 DB", "Compute Engine", "MongoDB", "정책 메타데이터 CRUD"],
                ["스케줄러", "Compute Engine", "Airflow 2.9.3", "DAG 자동 실행"],
                ["모니터링", "Compute Engine", "Prometheus + Grafana", "시스템 상태 감시"],
            ],
            caption="표 3. GCP 인프라 구성.",
        )
        self.p(
            "Cloud Run에 배포된 백엔드는 기동 시 GCS에서 FAISS 인덱스를 다운로드하여 "
            "메모리에 로드한다. scale-to-zero 설정으로 요청이 없을 때 비용이 발생하지 않으며, "
            "max 1 instance 제약으로 예산을 통제한다. Airflow는 별도 Compute Engine VM에서 "
            "운영되어 수집·인덱싱·평가 작업을 스케줄링한다."
        )

    # ── 4. 실험 및 결과 분석 ────────────────────────────
    def write_section4(self):
        self.h1("4. 실험 및 결과 분석")

        # 4.1
        self.h2("4.1. 실험 설정")

        self.h3("4.1.1. 평가 데이터셋 구성")
        self.p(
            "실험용 QA 데이터셋은 수집된 2,235건의 정책 문서에서 GPT-4o를 이용하여 100쌍의 "
            "질문-정답 쌍을 생성하였다(표 4). 각 QA 쌍은 질문(query), 정답(ground_truth), "
            "카테고리(category), 원본 정책 ID를 포함한다."
        )
        self.table(
            ["카테고리", "QA 쌍 수", "비율", "난이도 분포"],
            [
                ["교육(education)", "26", "26%", "-"],
                ["취업(employment)", "26", "26%", "-"],
                ["복지(welfare)", "24", "24%", "-"],
                ["주거(housing)", "24", "24%", "-"],
                ["합계", "100", "100%", "easy 46 / medium 36 / hard 18"],
            ],
            caption="표 4. QA 데이터셋 카테고리 분포.",
        )

        self.h3("4.1.2. 비교 대상 모델 및 전략")
        self.p(
            "독립변수는 검색 전략 4종(vector_only, bm25_only, hybrid, hybrid_rerank), "
            "생성 모델 5종(GPT-4o, GPT-4o-mini, Claude Sonnet, Gemini Flash, Llama3), "
            "RAG 유무 2종(RAG, No-RAG)이다. 총 실험 조건은 5모델 × RAG + 1 No-RAG = "
            "6조건이며, 각 조건당 100쌍을 평가하여 600건의 샘플을 생성하였다."
        )

        self.h3("4.1.3. 평가 지표 정의")
        self.table(
            ["지표", "유형", "범위", "측정 대상"],
            [
                ["Faithfulness", "RAGAS", "0~1", "답변의 컨텍스트 근거 비율"],
                ["Answer Relevancy", "RAGAS", "0~1", "답변의 질문 부합도"],
                ["Context Precision", "RAGAS", "0~1", "검색 컨텍스트 정밀도"],
                ["Context Recall", "RAGAS", "0~1", "검색 컨텍스트 재현율"],
                ["Citation Accuracy", "Judge", "1~5", "인용 정확도"],
                ["Completeness", "Judge", "1~5", "답변 완결성"],
                ["Readability", "Judge", "1~5", "답변 가독성"],
                ["Hallucination Score", "Safety", "0~1", "환각 수준 (낮을수록 안전)"],
            ],
            caption="표 5. 평가 지표 정의.",
        )

        # 4.2
        self.h2("4.2. 검색 전략별 성능 비교")
        self.p(
            "4가지 검색 전략의 성능을 RAGAS Context Precision과 Context Recall로 "
            "비교한다(표 6)."
        )
        self.table(
            ["전략", "Context Precision", "Context Recall", "Latency(s)"],
            [
                ["vector_only", "0.789 ± 0.308", "0.825 ± 0.363", "0.431"],
                ["bm25_only", "0.690 ± 0.356", "0.733 ± 0.438", "0.012"],
                ["hybrid (RRF)", "0.805 ± 0.312", "0.845 ± 0.351", "0.408"],
                ["hybrid_rerank", "0.795 ± 0.308", "0.855 ± 0.341", "0.449"],
            ],
            caption="표 6. 검색 전략별 성능 비교 (N=100).",
        )
        self.p(
            "Hybrid RRF가 Context Precision 0.805로 최고이며, hybrid_rerank가 Context "
            "Recall 0.855로 최고를 달성하였다(표 6). 이는 Cross-Encoder 리랭킹이 관련성 "
            "낮은 문서를 하위로 재배치하여 recall을 향상시키는 반면, precision은 RRF 단계에서 "
            "이미 최적화되기 때문으로 해석된다."
        )
        self.p(
            "BM25 단독 검색은 precision 0.690, recall 0.733으로 가장 낮았으나, 레이턴시 "
            "0.012초로 다른 전략(0.4~0.45초) 대비 약 35배 빠르다. 이는 실시간 응답이 중요한 "
            "환경에서의 트레이드오프를 시사한다. hybrid_rerank를 이후 생성 및 평가 실험의 "
            "고정 전략으로 채택하였으며, 이는 RAG 시스템에서 관련 문서를 최대한 포함하는 것이 "
            "답변 품질에 직접적 영향을 미치기 때문이다."
        )

        # 4.3
        self.h2("4.3. 멀티 LLM 응답 품질 비교")

        self.h3("4.3.1. RAGAS 정량 평가 결과")
        self.p("5개 모델의 RAG 응답에 대한 RAGAS 지표를 비교한다(표 7).")
        self.table(
            ["조건", "Faithfulness", "Answer Relevancy", "Context Precision", "Context Recall", "N"],
            [
                ["GPT-4o-mini (RAG)", "0.859 ± 0.289", "0.531 ± 0.232", "0.784 ± 0.318", "0.855 ± 0.341", "100"],
                ["GPT-4o (RAG)", "0.853 ± 0.291", "0.473 ± 0.281", "0.791 ± 0.312", "0.850 ± 0.350", "100"],
                ["Claude Sonnet (RAG)", "0.879 ± 0.138", "0.421 ± 0.235", "0.672 ± 0.397", "0.813 ± 0.377", "100"],
                ["Gemini Flash (RAG)", "0.826 ± 0.275", "0.485 ± 0.268", "0.782 ± 0.309", "0.855 ± 0.341", "100"],
                ["Llama3 (RAG)", "0.982 ± 0.058", "0.511 ± 0.178", "0.782 ± 0.308", "0.886 ± 0.295", "35"],
                ["GPT-4o-mini (No-RAG)", "N/A", "0.336 ± 0.320", "N/A", "N/A", "100"],
            ],
            caption="표 7. 모델별 RAGAS 지표.",
        )
        self.p(
            "Faithfulness 측면에서 Llama3가 0.982로 가장 높았으나, N=35로 표본이 적어 "
            "생존 편향 가능성을 고려해야 한다. 대규모 표본 모델 중에서는 Claude Sonnet이 "
            "0.879로 가장 높았으며, GPT-4o-mini가 0.859, GPT-4o가 0.853으로 유사한 수준이었다."
        )
        self.p(
            "Answer Relevancy에서는 GPT-4o-mini가 0.531로 최고를 기록하였으며, No-RAG "
            "대조군(0.336)과 비교하면 RAG 적용으로 0.195(58.0%) 향상되었다. Context "
            "Precision과 Context Recall은 검색 전략(hybrid_rerank)에 의존하므로 모델 간 "
            "차이가 미미하며, 이는 예상된 결과이다."
        )

        self.h3("4.3.2. LLM Judge 정성 평가 결과")
        self.p(
            "GPT-4o-mini를 Judge로 활용한 정성 평가 결과를 비교한다(표 8)."
        )
        self.table(
            ["조건", "Citation Acc.", "Completeness", "Readability", "Average", "N"],
            [
                ["GPT-4o-mini (RAG)", "4.59 ± 1.14", "4.55 ± 1.23", "4.98 ± 0.14", "4.71 ± 0.79", "100"],
                ["GPT-4o (RAG)", "4.39 ± 1.38", "4.29 ± 1.51", "4.88 ± 0.43", "4.52 ± 1.04", "100"],
                ["Gemini Flash (RAG)", "4.62 ± 1.02", "4.48 ± 1.16", "4.96 ± 0.28", "4.69 ± 0.76", "100"],
                ["Llama3 (RAG)", "4.94 ± 0.23", "4.93 ± 0.24", "4.99 ± 0.08", "4.95 ± 0.17", "35"],
                ["GPT-4o-mini (No-RAG)", "2.46 ± 1.47", "4.37 ± 1.16", "5.00 ± 0.05", "3.94 ± 0.71", "100"],
            ],
            caption="표 8. 모델별 LLM Judge 점수 (GPT-4o-mini Judge).",
            footer="Claude Sonnet은 유효 Judge 결과 24건으로 표에서 제외.",
        )
        self.p(
            "GPT-4o-mini RAG가 Average 4.71로 대규모 표본 모델 중 최고를 달성하였으며, "
            "Citation Accuracy 4.59, Completeness 4.55, Readability 4.98로 전 차원에서 "
            "균형 잡힌 성능을 보였다. Readability는 모든 모델에서 4.88~5.00으로 높아, "
            "한국어 답변의 가독성은 모델 간 차별화 요인이 아님을 시사한다."
        )
        self.p(
            "No-RAG 대조군은 Citation Accuracy 2.46으로 RAG 대비 2.13점 하락하여, "
            "RAG의 인용 정확도 기여가 가장 두드러졌다. 반면 Completeness(4.37)와 "
            "Readability(5.00)는 높은 수준을 유지하여, LLM의 일반 지식만으로도 "
            "완결성과 가독성은 확보 가능함을 보인다."
        )

        self.h3("4.3.3. 안전성(환각) 평가 결과")
        self.p(
            "DeepEval HallucinationMetric으로 측정한 환각 점수를 비교한다(표 9). "
            "Hallucination Score는 0.0이 환각 없음, 1.0이 완전한 환각을 의미하며 "
            "낮을수록 안전하다."
        )
        self.table(
            ["조건", "Hallucination Score", "N"],
            [
                ["Gemini Flash (RAG)", "0.338 ± 0.365", "100"],
                ["GPT-4o (RAG)", "0.350 ± 0.367", "100"],
                ["GPT-4o-mini (RAG)", "0.368 ± 0.380", "100"],
                ["Llama3 (RAG)", "0.457 ± 0.369", "35"],
                ["Claude Sonnet (RAG)", "0.533 ± 0.340", "100"],
                ["GPT-4o-mini (No-RAG)", "0.000 ± 0.000", "100"],
            ],
            caption="표 9. 모델별 환각 점수.",
        )
        self.p(
            "RAG 모델 중 Gemini Flash가 0.338로 가장 낮은 환각률을 보였으며, "
            "Claude Sonnet이 0.533으로 가장 높았다. Claude Sonnet의 높은 환각률은 "
            "컨텍스트를 넘어 추가 정보를 생성하는 경향에 기인한다."
        )
        self.p(
            "No-RAG 대조군의 환각 점수가 0.000인 것은 측정 한계에 해당한다. DeepEval의 "
            "환각 판정 기준이 컨텍스트 대비 사실 왜곡이므로, 컨텍스트 자체가 없으면 환각으로 "
            "판정할 근거가 부재하다. 이는 No-RAG 답변에 환각이 없다는 의미가 아니라, 측정 "
            "프레임워크의 구조적 한계이다."
        )

        # 4.4
        self.h2("4.4. LLM Judge 비용-성능 분석")
        self.p(
            "GPT-4o-mini와 Gemini 3.1 Pro를 Judge 모델로 사용하여 동일 600건을 평가한 결과, "
            "양측 모두 유효한 267건의 비교 쌍을 확보하였다(표 10)."
        )
        self.table(
            ["지표", "값", "해석"],
            [
                ["Kendall τ", "0.259", "약한 순위 상관 (통계적으로 유의, p<0.001)"],
                ["Spearman ρ", "0.268", "약한 순위 상관 (통계적으로 유의, p<0.001)"],
                ["MAE", "0.307", "평균 0.3점 차이 (5점 척도)"],
                ["Perfect Agreement", "86.9%", "정수 반올림 시 동일 비율"],
                ["Class Agreement", "90.3%", "3등급(low/mid/high) 분류 일치율"],
                ["입력 토큰 비용 비율", "13.3배", "GPT-4o-mini $0.15 vs Gemini 3.1 Pro $2.00 /1M"],
                ["출력 토큰 비용 비율", "20.0배", "GPT-4o-mini $0.60 vs Gemini 3.1 Pro $12.00 /1M"],
            ],
            caption=(
                "표 10. Judge 모델 비용-성능 비교 (GPT-4o-mini vs Gemini 3.1 Pro, N=267)."
                " 가격은 2026년 5월 공식 API 가격표 기준."
            ),
        )
        self.p(
            "Class Agreement Rate가 90.3%로, 3등급 분류(low/mid/high) 기준 10건 중 9건에서 "
            "두 Judge가 동일한 등급을 부여하였다. Readability에서 가장 높은 일치도(Class "
            "Agreement 98.5%, MAE 0.043)를 보인 반면, Completeness에서 가장 낮은 일치도"
            "(Class Agreement 87.3%, MAE 0.472)를 나타냈다. 가독성 평가는 모델 간 관점 차이가 "
            "적으나, 완결성 판단은 주관성이 개입하여 모델 간 시각 차이가 발생하는 것으로 해석된다."
        )
        self.p(
            "순위 상관(Kendall τ=0.259, Spearman ρ=0.268)이 낮은 것은 대부분의 점수가 "
            "4~5점 고점에 집중되어 변동성이 적기 때문이다. 점수 분포가 제한적인 상황에서 순위 "
            "상관 계수는 일치도를 과소평가하는 경향이 있으며, 따라서 Class Agreement가 더 "
            "적합한 평가 지표이다."
        )
        self.p(
            "비용 측면에서, 공식 API 가격표(2026년 5월 기준) GPT-4o-mini는 입력 $0.15/1M, "
            "출력 $0.60/1M이며 Gemini 3.1 Pro는 입력 $2.00/1M, 출력 $12.00/1M으로, "
            "입력 토큰 기준 약 13배, 출력 토큰 기준 약 20배의 가격 차이가 있다. "
            "90.3%의 평가 일관성을 이 비용 격차에서 달성한다는 점은 프로덕션 환경에서의 "
            "지속적 모니터링에 경량 Judge 활용이 경제적으로 타당함을 입증한다."
        )

        # 4.5
        self.h2("4.5. Position Bias 완화 효과 검증")
        self.p(
            "459건의 유효 Judge 결과에서 원본 순서와 셔플 순서 간 점수 차이를 분석하였다. "
            "Citation Accuracy에서 1점 이상 차이가 발생한 비율은 6.5%(30/459)로 가장 "
            "높았으며, Readability는 1.7%(8/459)로 가장 안정적이었다. Wilcoxon 부호순위 "
            "검정 결과, 모든 차원에서 p>0.05로 원본/셔플 간 체계적 편향은 통계적으로 유의하지 "
            "않았다."
        )
        self.p(
            "2회 평균의 분산 감소 효과는 2.79%로 미미하였다. 이는 GPT-4o-mini Judge가 "
            "컨텍스트 순서에 대해 원래 강건하여, 완화 기법의 추가 효과가 제한적임을 시사한다. "
            "그러나 Citation Accuracy에서 최대 4점 차이가 발생한 사례가 존재하므로, 2회 평균은 "
            "극단적 편향을 방지하는 안전망으로서의 가치가 있다."
        )
        self.fig_placeholder("그림 4. Position Bias 점수 차이 분포 (Citation Accuracy).")
        self.p(
            "3단계 평가의 상보성을 교차 상관 분석으로 검증하였다. RAGAS Faithfulness와 "
            "Judge Average 간 Spearman ρ=0.402, Faithfulness와 Safety Hallucination 간 "
            "ρ=0.261로, 3단계는 서로 다른 품질 차원을 측정함을 확인하였다(최대 ρ=0.502). "
            "탐지율 분석에서는 단독 평가 최고 29.7%(Safety) 대비 3단계 조합이 52.0%의 결함 "
            "탐지율을 달성하여, 조합의 가치가 실증되었다(그림 5). 특히 Judge+Safety 2단계 "
            "조합이 50.7%로, RAGAS의 추가 기여(+1.3%p)는 제한적이나 안전망 역할은 유효하다."
        )
        self.fig_placeholder("그림 5. 단독/2단계/3단계 평가 탐지율 비교.")

    # ── 5. 결론 ────────────────────────────────────────
    def write_section5(self):
        self.h1("5. 결론")

        self.h2("5.1. 연구 요약")
        self.p(
            "본 연구는 한국어 청년 정책 도메인에 특화된 하이브리드 RAG 질의응답 시스템과 "
            "비용 효율적 3단계 응답 품질 평가 파이프라인을 설계·구현하였다. 2,235건의 정책 "
            "문서를 대상으로 4가지 검색 전략과 5개 LLM의 성능을 체계적으로 비교하였으며, "
            "RAGAS v0.4, LLM-as-a-Judge, DeepEval을 결합한 평가 프레임워크의 상보성과 "
            "경량 Judge 모델의 비용 효율성을 실증하였다."
        )

        self.h2("5.2. 학술적 기여")
        self.p("본 연구의 학술적 기여는 다음과 같다.")
        self.item(
            "1. 한국어 정책 도메인에서 4가지 하이브리드 검색 전략(Vector, BM25, Hybrid RRF, "
            "Hybrid Rerank)의 체계적 성능 비교를 수행하여, hybrid_rerank가 Context Recall "
            "0.855로 최고 성능을 달성함을 확인하였다."
        )
        self.item(
            "2. RAGAS 정량 평가, LLM Judge 정성 평가, DeepEval 안전성 평가를 단일 "
            "파이프라인으로 통합한 3단계 상보적 프레임워크를 제안하고, 3단계 조합이 단독 "
            "평가 대비 22.3%p 높은 결함 탐지율(52.0%)을 달성함을 실증하였다."
        )
        self.item(
            "3. 파인튜닝 없이 경량 상용 모델(GPT-4o-mini)을 Judge로 활용하여, 고비용 모델"
            "(Gemini 3.1 Pro) 대비 90.3%의 평가 일치율을 입력 토큰 기준 약 13배 낮은 비용으로 달성 가능함을 "
            "검증하였다."
        )
        self.item(
            "4. RAG 컨텍스트 순서 기반 Position Bias 완화를 위한 2회 평균 기법이 극단적 "
            "점수 편향을 방지하는 안전망으로 기능함을 확인하였다."
        )

        self.h2("5.3. 한계 및 향후 연구")
        self.p(
            "본 연구의 한계와 향후 연구 방향은 다음과 같다. 첫째, 100쌍의 QA 데이터셋은 "
            "통계적 검정력이 제한적이며, 향후 500쌍 이상으로 확장하여 결과의 일반화 가능성을 "
            "높일 필요가 있다. 둘째, Judge 모델 비교가 GPT-4o-mini와 Gemini 3.1 Pro 2종에 "
            "한정되었으며, 한국어 특화 모델(KoGPT, EXAONE 등)과의 비교가 필요하다. 셋째, "
            "Cross-Encoder로 사용한 ms-marco-MiniLM-L-6-v2는 영어 기반 모델로, 한국어 "
            "의미론적 정밀도에 한계가 있을 수 있으며, 한국어 특화 Cross-Encoder 적용이 "
            "향후 과제이다."
        )

    # ── 감사의 글 ──────────────────────────────────────
    def write_acknowledgments(self):
        self.p(
            "Acknowledgments: 본 연구를 지도해 주신 장익범 교수님께 깊이 감사드립니다. "
            "또한 학부 과정 동안 학문적 성장을 도와주신 한국외국어대학교 컴퓨터공학부 "
            "교수님들과, 졸업 프로젝트를 함께한 동료들에게 감사의 뜻을 전합니다.",
            S_BACK,
        )

    # ── 참고문헌 ───────────────────────────────────────
    def write_references(self):
        self.h1("참고문헌")

        refs = [
            (
                "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., "
                "Küttler, H., Lewis, M., Yih, W., Rocktäschel, T., Riedel, S., Kiela, D. "
                "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. "
                "Advances in Neural Information Processing Systems (NeurIPS), 33, "
                "pp. 9459-9474, 2020"
            ),
            (
                "Es, S., James, J., Anke, L. E., Schockaert, S. "
                "RAGAS: Automated Evaluation of Retrieval Augmented Generation. "
                "Proceedings of the 18th Conference of the European Chapter of the "
                "Association for Computational Linguistics (EACL), pp. 150-163, 2024"
            ),
            (
                "Zheng, L., Chiang, W. L., Sheng, Y., Zhuang, S., Wu, Z., Zhuang, Y., "
                "Lin, Z., Li, Z., Li, D., Xing, E., Zhang, H., Gonzalez, J. E., Stoica, I. "
                "Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. "
                "Advances in Neural Information Processing Systems (NeurIPS), 36, 2023"
            ),
            (
                "Liu, Y., Iter, D., Xu, Y., Wang, S., Xu, R., Zhu, C. "
                "G-Eval: NLG Evaluation Using GPT-4 with Better Human Alignment. "
                "Proceedings of the 2023 Conference on Empirical Methods in Natural "
                "Language Processing (EMNLP), pp. 2511-2522, 2023"
            ),
            (
                "Wang, P., Li, L., Chen, L., Cai, Z., Zhu, D., Lin, B., Cao, Y., "
                "Liu, Q., Liu, T., Sui, Z. "
                "Large Language Models Are Not Fair Evaluators. "
                "Proceedings of the 62nd Annual Meeting of the Association for "
                "Computational Linguistics (ACL), 2024"
            ),
            (
                "Kim, S., Shin, J., Choi, Y., Min, J., Seo, M. "
                "Prometheus: Inducing Fine-Grained Evaluation Capability in "
                "Language Models. "
                "Proceedings of the 12th International Conference on Learning "
                "Representations (ICLR), 2024"
            ),
            (
                "Saad-Falcon, J., Khattab, O., Potts, C., Zaharia, M. "
                "ARES: An Automated Evaluation Framework for Retrieval-Augmented "
                "Generation Systems. "
                "Proceedings of the 2024 Conference of the North American Chapter "
                "of the Association for Computational Linguistics (NAACL), 2024"
            ),
            (
                "Cormack, G. V., Clarke, C. L., Buettcher, S. "
                "Reciprocal Rank Fusion Outperforms Condorcet and Individual "
                "Rank Learning Methods. "
                "Proceedings of the 32nd International ACM SIGIR Conference on "
                "Research and Development in Information Retrieval, pp. 758-759, 2009"
            ),
            (
                "Nogueira, R., Cho, K. "
                "Passage Re-Ranking with BERT. "
                "arXiv preprint arXiv:1901.04085, 2019"
            ),
            (
                "RAGAS Documentation. RAGAS v0.4 API Reference. "
                "https://docs.ragas.io/en/stable/, 2024"
            ),
        ]

        for i, ref in enumerate(refs, 1):
            self.p(f"[{i}] {ref}", S_REF)

    # ── main ───────────────────────────────────────────
    def write_all(self):
        self.write_cover()
        self.write_header()
        self.write_abstract()
        self.write_section1()
        self.write_section2()
        self.write_section3()
        self.write_section4()
        self.write_section5()
        self.write_acknowledgments()
        self.write_references()

    def save(self):
        self.doc.save(str(OUTPUT))


def main():
    writer = ThesisWriter()
    writer.write_all()
    writer.save()
    print(f"논문 생성 완료: {OUTPUT}")


if __name__ == "__main__":
    main()
